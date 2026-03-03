import streamlit as st
from docx import Document
from docx.shared import Inches
import io

# Configuração da Página
st.set_page_config(page_title="Vistoria Pro v2026", page_icon="📋", layout="centered")

# --- CSS DASHBOARD (ESTILO BASE44) ---
st.markdown("""
    <style>
    .stApp { background-color: #f1f5f9; }
    header {visibility: hidden;}
    .main-header {
        background-color: #1e293b; padding: 15px; color: white;
        text-align: center; font-size: 1.4rem; font-weight: bold;
        border-radius: 0 0 12px 12px; margin-bottom: 20px;
    }
    .stTabs [data-baseweb="tab-list"] { gap: 5px; }
    .stTabs [data-baseweb="tab"] { 
        background-color: #e2e8f0; border-radius: 8px 8px 0 0; padding: 8px 15px; font-weight: 600;
    }
    .stTabs [aria-selected="true"] { background-color: #2563eb !important; color: white !important; }
    .stCheckbox { margin-bottom: -15px; }
    </style>
    <div class="main-header">📋 Sistema de Vistoria Profissional</div>
    """, unsafe_allow_html=True)

# --- MEMÓRIA ---
if 'etapa' not in st.session_state: st.session_state.etapa = 1
if 'historico' not in st.session_state: st.session_state.historico = []
if 'fotos_db' not in st.session_state: st.session_state.fotos_db = {}

# --- FUNÇÕES AUXILIARES ---
def resetar():
    for key in list(st.session_state.keys()):
        if key not in ['historico']: del st.session_state[key]
    st.session_state.etapa = 1
    st.rerun()

# --- ETAPA 1: IDENTIFICAÇÃO ---
if st.session_state.etapa == 1:
    st.subheader("📍 Identificação")
    with st.expander("Dados Gerais", expanded=True):
        end = st.text_input("Endereço do Imóvel", key="main_end")
        c1, c2 = st.columns(2)
        data_v = c1.date_input("Data da Vistoria")
        tipo_v = c2.selectbox("Tipo", ["Entrada", "Saída", "Transferência"])
        if st.button("Configurar Cômodos ➡️"):
            if end:
                st.session_state.info_geral = {"end": end, "data": str(data_v), "tipo": tipo_v}
                st.session_state.etapa = 2
                st.rerun()

# --- ETAPA 2: CÔMODOS ---
elif st.session_state.etapa == 2:
    st.subheader("🏠 O que tem no imóvel?")
    opcoes = ["Sala", "Cozinha", "Banheiro Social", "Área de Serviço", "Corredor", "Varanda", "Garagem"]
    selecionados = [item for item in opcoes if st.checkbox(item, key=f"sel_{item}")]
    c1, c2 = st.columns(2)
    q = c1.number_input("Quartos Simples", 0, 10, 1)
    s = c2.number_input("Suítes", 0, 10, 0)
    if st.button("Iniciar Inspeção Detalhada 📝"):
        lista = selecionados + [f"Quarto {i+1}" for i in range(q)]
        for i in range(s):
            lista.append(f"Suíte {i+1}"); lista.append(f"Banheiro Suíte {i+1}")
        st.session_state.comodos = lista
        st.session_state.etapa = 3
        st.rerun()

# --- ETAPA 3: INSPEÇÃO TÉCNICA ---
elif st.session_state.etapa == 3:
    st.info(f"📍 {st.session_state.info_geral['end']}")
    tabs = st.tabs(st.session_state.comodos)
    
    # LISTAS TÉCNICAS
    ESTADOS = ["Bom estado", "Novo", "Usado (Limpo)", "Manchado", "Riscado", "Trincado", "Faltante"]
    CORES = ["Branca", "Gelo", "Cinza", "Bege", "Madeira", "Preta", "Inox", "Verde", "Azul"]
    MAT_PISO = ["Porcelanato", "Laminado", "Vinílico", "Cerâmico", "Taco/Madeira", "Ardósia", "Cimento Queimado", "Carpete"]
    MAT_BANCADA = ["Granito", "Mármore", "Inox", "MDF Revestido", "Quartzo"]

    for i, nome_comodo in enumerate(st.session_state.comodos):
        with tabs[i]:
            kb = f"{nome_comodo}_{i}"
            
            # --- PISO E ACABAMENTOS ---
            with st.expander("🏗️ Revestimentos (Piso/Parede/Teto)", expanded=True):
                # Piso
                st.markdown("**Piso**")
                c1, c2, c3 = st.columns(3)
                p_mat = c1.selectbox("Material", MAT_PISO, key=f"pm_{kb}")
                p_cor = c2.selectbox("Cor", CORES, key=f"pc_{kb}")
                p_est = c3.selectbox("Estado", ESTADOS, key=f"pe_{kb}")
                
                # Rodapé
                if st.checkbox("Tem Rodapé?", value=True, key=f"tr_{kb}"):
                    c1, c2 = st.columns(2)
                    r_mat = c1.selectbox("Material Rodapé", ["Mesmo do piso", "Madeira/MDF", "PVC"], key=f"rm_{kb}")
                    r_est = c2.selectbox("Estado Rodapé", ESTADOS, key=f"re_{kb}")
                
                # Paredes e Teto
                st.markdown("---")
                c1, c2 = st.columns(2)
                pa_est = c1.selectbox("Pintura Paredes", ESTADOS, key=f"pae_{kb}")
                te_est = c2.selectbox("Estado Teto", ESTADOS, key=f"tee_{kb}")
                st.checkbox("Moldura de Gesso?", key=f"m_g_{kb}")
                
                f_rev = st.file_uploader("📷 Foto Revestimentos", type=['jpg', 'png'], key=f"f_rev_{kb}")
                if f_rev: st.session_state.fotos_db[f"f_rev_{kb}"] = f_rev

            # --- ABERTURAS ---
            with st.expander("🚪 Aberturas (Porta/Janela)"):
                st.markdown("**Porta**")
                c1, c2, c3 = st.columns(3)
                po_fol = c1.selectbox("Folha", ESTADOS, key=f"poe_{kb}")
                po_bat = c2.selectbox("Batente", ESTADOS, key=f"pob_{kb}")
                po_mac = c3.selectbox("Maçaneta", ESTADOS, key=f"pom_{kb}")
                st.number_input("Qtd Chaves", 0, 5, key=f"ch_{kb}")
                
                st.markdown("**Janela**")
                c1, c2 = st.columns(2)
                ja_mat = c1.selectbox("Material Janela", ["Alumínio", "Madeira", "PVC", "Ferro", "Blindex"], key=f"jam_{kb}")
                ja_est = c2.selectbox("Estado Vidros/Trincos", ESTADOS, key=f"jae_{kb}")
                
                f_ab = st.file_uploader("📷 Foto Aberturas", type=['jpg', 'png'], key=f"f_ab_{kb}")
                if f_ab: st.session_state.fotos_db[f_ab_{kb}] = f_ab

            # --- ELÉTRICA ---
            with st.expander("💡 Elétrica e Iluminação"):
                c1, c2, c3 = st.columns(3)
                tom_q = c1.number_input("Tomadas", 0, 30, key=f"tq_{kb}")
                int_q = c2.number_input("Interruptores", 0, 20, key=f"iq_{kb}")
                el_est = c3.selectbox("Estado Espelhos", ESTADOS, key=f"ele_{kb}")
                
                st.markdown("**Iluminação**")
                c1, c2 = st.columns(2)
                l_tip = c1.selectbox("Tipo", ["Plafon", "Spot LED", "Lâmpada Simples", "Lustre"], key=f"lti_{kb}")
                l_sta = c2.radio("Status", ["OK", "Queimada", "Faltante"], key=f"l_st_{kb}", horizontal=True)
                
                f_el = st.file_uploader("📷 Foto Elétrica", type=['jpg', 'png'], key=f"f_el_{kb}")
                if f_el: st.session_state.fotos_db[f"f_el_{kb}"] = f_el

            # --- ÁREAS MOLHADAS / MÓVEIS (Cozinha/Banheiro) ---
            if "Cozinha" in nome_comodo or "Banheiro" in nome_comodo or "Serviço" in nome_comodo:
                with st.expander("🚰 Hidráulica e Mobiliário"):
                    if st.checkbox("Tem Bancada/Pia?", key=f"t_ba_{kb}"):
                        c1, c2 = st.columns(2)
                        c1.selectbox("Material Bancada", MAT_BANCADA, key=f"bam_{kb}")
                        c2.selectbox("Estado Bancada", ESTADOS, key=f"bae_{kb}")
                    
                    if st.checkbox("Tem Torneira/Metais?", key=f"t_me_{kb}"):
                        c1, c2 = st.columns(2)
                        c1.selectbox("Estado Torneira", ESTADOS, key=f"tme_{kb}")
                        c2.checkbox("Testado/Funcionando?", key=f"tm_f_{kb}")
                    
                    if st.checkbox("Tem Armários?", key=f"t_ar_{kb}"):
                        st.selectbox("Estado Portas/Gavetas", ESTADOS, key=f"are_{kb}")
                    
                    if "Banheiro" in nome_comodo:
                        st.selectbox("Vaso Sanitário/Assento", ESTADOS, key=f"vse_{kb}")
                        st.selectbox("Box/Chuveiro", ESTADOS, key=f"bce_{kb}")

                    f_hid = st.file_uploader("📷 Foto Hidráulica/Móveis", type=['jpg', 'png'], key=f"f_hid_{kb}")
                    if f_hid: st.session_state.fotos_db[f"f_hid_{kb}"] = f_hid

            st.text_area("Notas Adicionais", key=f"obs_{kb}", placeholder="Relate detalhes como vazamentos, manchas de infiltração, etc.")

    # --- GERAÇÃO FINAL ---
    if st.button("🚀 GERAR RELATÓRIO PROFISSIONAL"):
        doc = Document()
        doc.add_heading('LAUDO TÉCNICO DE VISTORIA', 0)
        doc.add_paragraph(f"IMÓVEL: {st.session_state.info_geral['end']}")
        doc.add_paragraph(f"DATA: {st.session_state.info_geral['data']} | TIPO: {st.session_state.info_geral['tipo']}")

        for i, nome in enumerate(st.session_state.comodos):
            kb = f"{nome}_{i}"
            doc.add_heading(nome.upper(), level=1)
            
            # Texto Técnico Gerado
            txt = f"Inspeção realizada no cômodo {nome}. "
            txt += f"Piso em {st.session_state.get(f'pm_{kb}')} na cor {st.session_state.get(f'pc_{kb}')} encontra-se {st.session_state.get(f'pe_{kb}')}. "
            
            if st.session_state.get(f'tr_{kb}'):
                txt += f"Rodapé em {st.session_state.get(f'rm_{kb}')} em {st.session_state.get(f're_{kb}')}. "
            
            txt += f"Paredes com pintura em {st.session_state.get(f'pae_{kb}')} e teto em {st.session_state.get(f'tee_{kb}')}. "
            
            txt += f"Porta: folha {st.session_state.get(f'poe_{kb}')}, batente {st.session_state.get(f'pob_{kb}')} e maçaneta {st.session_state.get(f'pom_{kb}')}. "
            txt += f"Janela em {st.session_state.get(f'jam_{kb}')} com vidros e trincos em {st.session_state.get(f'jae_{kb}')}. "
            
            txt += f"Elétrica composta por {st.session_state.get(f'tq_{kb}')} tomadas e {st.session_state.get(f'iq_{kb}')} interruptores, com espelhos em {st.session_state.get(f'ele_{kb}')}. "
            txt += f"Iluminação {st.session_state.get(f'lti_{kb}')} está {st.session_state.get(f'l_st_{kb}')}. "
            
            if st.session_state.get(f't_ba_{kb}'):
                txt += f"Bancada em {st.session_state.get(f'bam_{kb}')} em {st.session_state.get(f'bae_{kb}')}. "
            
            doc.add_paragraph(txt)
            if st.session_state.get(f'ob_{kb}'):
                doc.add_paragraph(f"OBS: {st.session_state.get(f'ob_{kb}')}")

            # FOTOS 2x2
            fotos_locais = []
            for f_key in [f"f_rev_{kb}", f"f_ab_{kb}", f"f_el_{kb}", f"f_hid_{kb}"]:
                if f_key in st.session_state.fotos_db:
                    fotos_locais.append(st.session_state.fotos_db[f_key])

            if fotos_locais:
                table = doc.add_table(rows=(len(fotos_locais)+1)//2, cols=2)
                for idx, photo in enumerate(fotos_locais):
                    img_stream = io.BytesIO(photo.getvalue())
                    cell = table.rows[idx // 2].cells[idx % 2]
                    run = cell.paragraphs[0].add_run()
                    run.add_picture(img_stream, width=Inches(3.0))

        target = io.BytesIO()
        doc.save(target)
        st.download_button("📥 BAIXAR LAUDO FINALIZADO", target.getvalue(), "vistoria_pro.docx")
